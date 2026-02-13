"""
Script pour générer le rapport final Rakuten en format DOCX
avec un design professionnel et des marges correctes.
"""

import subprocess
import sys

# Install python-docx if not available
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C41E3A')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def create_rapport():
    # Create document
    doc = Document()

    # Set page margins (2.5cm all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Define styles
    styles = doc.styles

    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Georgia'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(6)

    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Georgia'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(10)

    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Georgia'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x1B, 0x2B, 0x4B)
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(8)

    # Normal text style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.space_after = Pt(8)

    # ==================== COVER PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()

    # Logo R
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("R")
    run.font.name = 'Georgia'
    run.font.size = Pt(72)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    doc.add_paragraph()

    # Institution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DATASCIENTEST × MINES PARIS - PSL")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Formation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Machine Learning Engineer (BMLE) — Promotion Octobre 2025")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2B, 0x4B)

    doc.add_paragraph()
    add_horizontal_line(doc)
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Classification Multimodale\n")
    run.font.name = 'Georgia'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    run = p.add_run("de Produits E-Commerce")
    run.font.name = 'Georgia'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Projet Rakuten France — Challenge de Classification Automatique\n")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    run = p.add_run("Approche Hybride Texte + Image avec Voting System")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1B, 0x2B, 0x4B)

    doc.add_paragraph()
    add_horizontal_line(doc)
    doc.add_paragraph()

    # Team
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ÉQUIPE PROJET")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    for author in ["Johan Frachon", "Liviu Andronic", "Hery Mickael Ralaimanantsoa", "Oussama Akir"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()

    # Mentor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mentor : ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run = p.add_run("Antoine")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()

    # Metrics table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = table.rows[0].cells

    metrics = [
        ("84 916", "Produits"),
        ("27", "Catégories"),
        ("~79%", "Accuracy Image"),
        ("83%", "Accuracy Texte")
    ]

    for i, (value, label) in enumerate(metrics):
        cells[i].width = Cm(3.5)
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value + "\n")
        run.font.name = 'Georgia'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)
        run = p.add_run(label)
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Février 2025")
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_page_break()

    # ==================== RÉSUMÉ EXÉCUTIF ====================
    doc.add_heading("Résumé Exécutif", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Ce rapport présente notre solution de ")
    run = p.add_run("classification automatique multimodale")
    run.bold = True
    run = p.add_run(" développée dans le cadre du challenge Rakuten France. L'objectif était de classifier automatiquement des produits e-commerce parmi ")
    run = p.add_run("27 catégories")
    run.bold = True
    run = p.add_run(" en utilisant à la fois les données textuelles (titre et description) et visuelles (images produits).")

    p = doc.add_paragraph()
    run = p.add_run("Notre approche hybride combine un ")
    run = p.add_run("classifieur textuel LinearSVC")
    run.bold = True
    run = p.add_run(" (TF-IDF word+char, accuracy 83%) et un ")
    run = p.add_run("Voting System d'images")
    run.bold = True
    run = p.add_run(" fusionnant trois architectures complémentaires : DINOv3 (Vision Transformer), XGBoost sur features ResNet, et EfficientNet-B0. Ce système de vote atteint ")
    run = p.add_run("~79% d'accuracy")
    run.bold = True
    run = p.add_run(" sur les images seules.")

    p = doc.add_paragraph()
    p.add_run("La fusion tardive (Late Fusion) des deux modalités avec pondération optimisée permet d'atteindre des performances robustes sur l'ensemble des 27 catégories, y compris les classes minoritaires grâce aux stratégies d'oversampling et de class weighting.")

    # Keywords
    doc.add_heading("Mots-clés", level=2)
    keywords = ["Classification Multimodale", "Transfer Learning", "Voting Classifier", "TF-IDF", "Vision Transformer", "E-commerce", "Deep Learning"]
    p = doc.add_paragraph()
    for i, kw in enumerate(keywords):
        run = p.add_run(kw)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        if i < len(keywords) - 1:
            p.add_run(" • ")

    # Results table
    doc.add_heading("Résultats Clés", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    results = [
        ("Accuracy Image (Voting)", "~79%"),
        ("Accuracy Texte (LinearSVC)", "83%"),
        ("Fusion Multimodale", "F1~0.85"),
        ("Catégories Classifiées", "27")
    ]

    for i, (label, value) in enumerate(results):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ==================== TABLE DES MATIÈRES ====================
    doc.add_heading("Table des Matières", level=1)

    toc_items = [
        ("PARTIE I : CONTEXTE ET DONNÉES", [
            "1.1 Le Challenge Rakuten France",
            "1.2 Description du Dataset",
            "1.3 Analyse Exploratoire (EDA)"
        ]),
        ("PARTIE II : PREPROCESSING & FEATURE ENGINEERING", [
            "2.1 Pipeline de Prétraitement Texte",
            "2.2 Pipeline de Prétraitement Image",
            "2.3 Gestion du Déséquilibre des Classes"
        ]),
        ("PARTIE III : MODÉLISATION TEXTE", [
            "3.1 Benchmark des Classifieurs",
            "3.2 Optimisation LinearSVC",
            "3.3 Résultats Détaillés par Classe"
        ]),
        ("PARTIE IV : MODÉLISATION IMAGE", [
            "4.1 Stratégie Transfer Learning",
            "4.2 Benchmark Machine Learning",
            "4.3 Approche Deep Learning",
            "4.4 Architectures Avancées (DINOv3, EfficientNet)",
            "4.5 Voting System - Modèle Final",
            "4.6 Tests de Robustesse"
        ]),
        ("PARTIE V : FUSION MULTIMODALE", [
            "5.1 Stratégie de Fusion Tardive",
            "5.2 Optimisation des Poids",
            "5.3 Résultats Combinés"
        ]),
        ("PARTIE VI : APPLICATION STREAMLIT", [
            "6.1 Architecture de l'Application",
            "6.2 Fonctionnalités et Interface"
        ]),
        ("PARTIE VII : CONCLUSION ET PERSPECTIVES", [
            "7.1 Bilan du Projet",
            "7.2 Limites et Difficultés",
            "7.3 Perspectives d'Amélioration"
        ])
    ]

    for part_title, sections in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(part_title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

        for section in sections:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(section)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_page_break()

    # ==================== PARTIE I ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARTIE I")
    run.font.name = 'Georgia'
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Contexte et Données")
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Challenge Rakuten France et analyse du dataset")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # Section 1.1
    doc.add_heading("1.1 Le Challenge Rakuten France", level=1)

    p = doc.add_paragraph()
    p.add_run("Rakuten, géant mondial du e-commerce, fait face à un défi classique des marketplaces : la catégorisation automatique des produits mis en ligne par des vendeurs tiers. Une mauvaise catégorisation entraîne une mauvaise expérience de recherche et une perte de revenus significative.")

    doc.add_heading("Objectif du Challenge", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Développer un modèle de classification multimodale capable de prédire le code catégorie (prdtypecode) d'un produit en utilisant simultanément :")

    p = doc.add_paragraph("• Le Texte : Désignation (titre) et description du produit", style='List Bullet')
    p = doc.add_paragraph("• L'Image : Visuel du produit fourni par le vendeur", style='List Bullet')

    doc.add_heading("Enjeux Business", level=2)
    p = doc.add_paragraph("• Expérience utilisateur : Navigation facilitée", style='List Bullet')
    p = doc.add_paragraph("• Recherche produit : Résultats pertinents", style='List Bullet')
    p = doc.add_paragraph("• Conversion : Réduction du taux de rebond", style='List Bullet')
    p = doc.add_paragraph("• Scalabilité : Millions de produits/jour", style='List Bullet')

    doc.add_heading("Métrique d'Évaluation", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Conformément aux règles du challenge, la métrique principale est le ")
    run = p.add_run("F1-Score weighted")
    run.bold = True
    run = p.add_run(", qui prend en compte le déséquilibre des classes en pondérant chaque classe par son support.")

    # Section 1.2
    doc.add_heading("1.2 Description du Dataset", level=1)

    # Dataset stats table
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ["Images Train", "Images Test", "Catégories", "Taille Images"]
    values = ["84 916", "13 812", "27", "500×500 px"]

    for i, (header, value) in enumerate(zip(headers, values)):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'F3F4F6')
        table.rows[1].cells[i].text = value
        table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_heading("Structure des Données", level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    data_structure = [
        ["Variable", "Type", "Description", "Complétude"],
        ["designation", "String", "Titre du produit", "100%"],
        ["description", "String", "Description longue (HTML/brut)", "65%"],
        ["productid", "Integer", "Identifiant unique produit", "100%"],
        ["imageid", "Integer", "Identifiant image associée", "100%"],
        ["prdtypecode", "Integer", "Code catégorie (target)", "100%"]
    ]

    for i, row_data in enumerate(data_structure):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], 'C41E3A')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("⚠️ Valeurs Manquantes Critiques : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x94, 0x0B)
    run = p.add_run("35% des descriptions sont manquantes (NaN). Cette contrainte structurelle nous oblige à concevoir un pipeline robuste qui ne dépend pas uniquement du champ description.")

    # Section 1.3
    doc.add_heading("1.3 Analyse Exploratoire (EDA)", level=1)

    doc.add_heading("Déséquilibre des Classes", level=2)
    p = doc.add_paragraph()
    p.add_run("L'analyse de la distribution des catégories révèle un déséquilibre significatif constituant l'un des défis majeurs du projet.")

    # Class imbalance table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    imbalance_data = [
        ["Métrique", "Valeur", "Classe", "Effectif"],
        ["Ratio Max/Min", "13.4×", "Majoritaire (Piscines)", "10 217"],
        ["Moyenne/Classe", "3 145", "Minoritaire (Consoles)", "761"],
        ["", "", "", ""]
    ]

    for i, row_data in enumerate(imbalance_data[:3]):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], '1B2B4B')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_page_break()

    # ==================== PARTIE II ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARTIE II")
    run.font.name = 'Georgia'
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Preprocessing & Feature Engineering")
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_page_break()

    # Section 2.1
    doc.add_heading("2.1 Pipeline de Prétraitement Texte", level=1)

    p = doc.add_paragraph()
    p.add_run("Le preprocessing textuel vise à transformer les descriptions produits brutes en vecteurs numériques exploitables par les algorithmes de classification.")

    doc.add_heading("Étapes de Nettoyage", level=2)
    p = doc.add_paragraph()
    run = p.add_run("TEXTE BRUT → NETTOYAGE → IMPUTATION → CONCATÉNATION → TF-IDF")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    doc.add_heading("Vectorisation TF-IDF", level=2)
    p = doc.add_paragraph()
    p.add_run("Nous avons opté pour une approche FeatureUnion combinant deux vectoriseurs complémentaires :")

    p = doc.add_paragraph()
    run = p.add_run("TF-IDF Word (N-grams) : ")
    run.bold = True
    run = p.add_run("Analyzer word, N-grams (1,2), Max features 120 000, Sublinear TF True")

    p = doc.add_paragraph()
    run = p.add_run("TF-IDF Char (N-grams) : ")
    run.bold = True
    run = p.add_run("Analyzer char_wb, N-grams (3,5), Max features 160 000, robuste aux fautes d'orthographe")

    # Section 2.2
    doc.add_heading("2.2 Pipeline de Prétraitement Image", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Le preprocessing image utilise le ")
    run = p.add_run("Transfer Learning")
    run.bold = True
    run = p.add_run(" avec EfficientNet-B0 pour extraire des features visuelles compactes et sémantiquement riches.")

    p = doc.add_paragraph()
    run = p.add_run("IMAGE 500×500 → RESIZE 224×224 → TO TENSOR → NORMALIZE → CNN → FEATURES 1×1280")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x2B, 0x4B)

    doc.add_page_break()

    # ==================== PARTIE III ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARTIE III")
    run.font.name = 'Georgia'
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Modélisation Texte")
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_page_break()

    # Section 3.1
    doc.add_heading("3.1 Benchmark des Classifieurs Texte", level=1)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    text_benchmark = [
        ["Modèle", "F1-Score", "Temps", "Verdict"],
        ["Logistic Regression", "0.79", "~2 min", "Standard"],
        ["Multinomial NB", "0.76", "~30 sec", "Rapide"],
        ["LinearSVC (C=0.5)", "0.83", "~3 min", "✓ Champion"],
        ["Random Forest", "0.68", "~15 min", "Lent"]
    ]

    for i, row_data in enumerate(text_benchmark):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], 'C41E3A')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 3:  # Winner row
                set_cell_shading(table.rows[i].cells[j], 'D1FAE5')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("✅ Champion Texte : LinearSVC ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run = p.add_run("avec class_weight='balanced' et C=0.5 optimisé par GridSearch atteint 83% d'accuracy.")

    doc.add_page_break()

    # ==================== PARTIE IV ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARTIE IV")
    run.font.name = 'Georgia'
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Modélisation Image")
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Du Transfer Learning au Voting System à ~79%")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # Section 4.2
    doc.add_heading("4.2 Benchmark Machine Learning", level=1)

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    ml_benchmark = [
        ["Modèle", "F1-Score", "Temps", "Hardware", "Verdict"],
        ["Random Forest", "0.71", "~30 min", "CPU", "Baseline"],
        ["XGBoost GPU", "0.72", "~10 min", "GPU", "Standard"],
        ["LightGBM", "0.71", "~5 min", "CPU", "Rapide"],
        ["XGBoost Heavy", "0.765", "6 heures", "CPU 128GB", "Force brute"]
    ]

    for i, row_data in enumerate(ml_benchmark):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], '1B2B4B')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("⚠️ Plafond de Performance ML : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD4, 0x94, 0x0B)
    run = p.add_run("Les modèles ML classiques plafonnent autour de 0.72-0.76 F1. Il faut passer au Deep Learning pour franchir ce plafond.")

    # Section 4.3
    doc.add_heading("4.3 Approche Deep Learning", level=1)

    p = doc.add_paragraph()
    p.add_run("Le passage aux réseaux de neurones denses (MLP) via PyTorch a provoqué une rupture dans les performances.")

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    dl_benchmark = [
        ["Optimizer", "Activation", "Dropout", "F1-Score", "Temps"],
        ["Adam", "GELU", "0.2", "0.9141", "58 sec"],
        ["Adam", "ReLU", "0.3", "0.9023", "55 sec"],
        ["RMSProp", "GELU", "0.2", "0.8956", "62 sec"],
        ["SGD", "ReLU", "0.3", "0.8734", "70 sec"]
    ]

    for i, row_data in enumerate(dl_benchmark):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], 'C41E3A')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 1:  # Winner row
                set_cell_shading(table.rows[i].cells[j], 'D1FAE5')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("🚀 Rupture de Performance : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run = p.add_run("Le Deep Learning brise le plafond de 76% pour atteindre 91%+. L'accélération GPU réduit le temps de 6 heures à moins de 60 secondes.")

    # Section 4.5
    doc.add_heading("4.5 Voting System - Modèle Final", level=1)

    p = doc.add_paragraph()
    p.add_run("Plutôt que de miser sur un seul modèle, nous avons construit un Voting Classifier exploitant la complémentarité des architectures.")

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    voting_data = [
        ["Modèle", "Score Solo", "Poids", "Rôle"],
        ["DINOv3", "79.1%", "0.40", "Vision globale, très confiant"],
        ["XGBoost/ResNet", "80.1%", "0.35", "Champion ML, textures"],
        ["EfficientNet", "~75%", "0.25", "Stabilisateur, détails fins"]
    ]

    for i, row_data in enumerate(voting_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], '1B2B4B')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Big result
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RÉSULTAT VOTING SYSTEM : ~79% ACCURACY")
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    doc.add_page_break()

    # ==================== PARTIE V ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARTIE V")
    run.font.name = 'Georgia'
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fusion Multimodale")
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_page_break()

    # Section 5.1
    doc.add_heading("5.1 Stratégie de Fusion Tardive (Late Fusion)", level=1)

    p = doc.add_paragraph()
    p.add_run("La fusion multimodale combine les prédictions des modèles texte et image pour exploiter la complémentarité des deux sources d'information.")

    doc.add_heading("Formule Late Fusion", level=2)
    p = doc.add_paragraph()
    run = p.add_run("P_final(classe) = α × P_image(classe) + (1-α) × P_texte(classe)")
    run.bold = True
    run.font.name = 'Consolas'

    p = doc.add_paragraph()
    run = p.add_run("Avec α = 0.6 (poids image) et (1-α) = 0.4 (poids texte)")
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Section 5.3
    doc.add_heading("5.3 Résultats Combinés", level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    fusion_data = [
        ["Configuration", "Accuracy", "F1 Weighted", "Avantage"],
        ["Texte seul (LinearSVC)", "83%", "0.83", "Rapide, interprétable"],
        ["Image seule (Voting)", "~79%", "~0.79", "Diversité architecturale"],
        ["Fusion Multimodale", "~85%", "~0.85", "Robustesse maximale"]
    ]

    for i, row_data in enumerate(fusion_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], 'C41E3A')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 3:  # Winner row
                set_cell_shading(table.rows[i].cells[j], 'D1FAE5')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("✅ Gain de la Fusion : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run = p.add_run("+2 points par rapport au meilleur modèle seul. Plus important : amélioration de la robustesse sur les cas difficiles.")

    doc.add_page_break()

    # ==================== PARTIE VII ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARTIE VII")
    run.font.name = 'Georgia'
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Conclusion et Perspectives")
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_page_break()

    # Section 7.1
    doc.add_heading("7.1 Bilan du Projet", level=1)

    p = doc.add_paragraph()
    p.add_run("Ce projet a permis de développer une solution complète de classification multimodale pour le challenge Rakuten France.")

    doc.add_heading("Contributions Principales", level=2)
    p = doc.add_paragraph("• Méthodologie : Pipeline complet de preprocessing texte et image avec gestion du déséquilibre de classes (ratio 1:13)", style='List Bullet')
    p = doc.add_paragraph("• Performance : Voting System innovant combinant 3 architectures pour atteindre ~79% sur les images", style='List Bullet')
    p = doc.add_paragraph("• Application : Interface Streamlit multimodale fonctionnelle avec visualisations et explicabilité", style='List Bullet')

    # Section 7.3
    doc.add_heading("7.3 Perspectives d'Amélioration", level=1)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    perspectives_data = [
        ["Piste", "Description", "Gain Estimé", "Effort"],
        ["OCR sur Images", "Lire le texte présent sur les images", "+3-5%", "Moyen"],
        ["CamemBERT/BERT", "Remplacer TF-IDF par embeddings contextuels", "+5-8%", "Élevé"],
        ["Early Fusion", "Concaténer features avant classification", "+2-3%", "Moyen"],
        ["Fine-tuning DINOv3", "Réentraîner sur le dataset", "+3-5%", "Très élevé"],
        ["Déploiement Cloud", "API REST sur AWS/GCP", "Scalabilité", "Moyen"]
    ]

    for i, row_data in enumerate(perspectives_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(table.rows[i].cells[j], '1B2B4B')
                table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("🎯 Priorité Recommandée : OCR ")
    run.bold = True
    run = p.add_run("- L'intégration d'un module OCR permettrait de 'lire' le texte présent sur les images, créant des features textuelles artificielles qui renforceraient la robustesse.")

    # ==================== FOOTER ====================
    doc.add_page_break()

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Projet Rakuten — Classification Multimodale")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Machine Learning Engineer — DataScientest × Mines Paris - PSL")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Février 2025")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Save document
    output_path = r"D:\datascientest\workspace\OCT25_BMLE_RAKUTEN_WS\repo\OCT25_BMLE_RAKUTEN\reports\RAPPORT_FINAL_RAKUTEN.docx"
    doc.save(output_path)
    print(f"[OK] Rapport DOCX genere avec succes : {output_path}")
    return output_path

if __name__ == "__main__":
    create_rapport()
